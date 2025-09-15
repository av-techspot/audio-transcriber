import sys, os
import argparse
from pathlib import Path
import whisper
from docx import Document

def hms(t):
	mm, ss = divmod(int(t), 60)
	hh, mm = divmod(mm, 60)
	return f"{hh:02d}:{mm:02d}:{ss:02d}"

def transcribe_to_docx(audio_path, out_docx, model_name="small", models_dir=None):
	try:
		print("Загружаю модель Whisper...")
		model = whisper.load_model(model_name, download_root=models_dir)
		print("Транскрибирую аудио...")
		# Проверяем FFmpeg
		import subprocess
		try:
			subprocess.run(["ffmpeg", "-version"], capture_output=True, check=True)
			print("FFmpeg найден")
		except (subprocess.CalledProcessError, FileNotFoundError):
			print("ОШИБКА: FFmpeg не найден в PATH!")
			print("Установите FFmpeg:")
			print("1. Через Chocolatey: choco install ffmpeg")
			print("2. Или скачайте с https://ffmpeg.org/download.html")
			print("3. Добавьте путь к ffmpeg.exe в переменную PATH")
			return False
		
		res = model.transcribe(audio_path, word_timestamps=False, verbose=False)
		print("Создаю DOCX файл...")
		doc = Document()
		doc.add_heading(os.path.basename(audio_path), 0)
		for s in res.get("segments", []):
			start, end, text = s.get("start",0), s.get("end",0), s.get("text","")
			doc.add_paragraph(f"[{hms(start)} - {hms(end)}] {text.strip()}")
		doc.save(out_docx)
		print(f"Готово! Файл сохранен: {out_docx}")
	except Exception as e:
		print(f"Ошибка: {e}")
		print(f"Тип ошибки: {type(e).__name__}")
		return False
	return True

if __name__ == "__main__":
	parser = argparse.ArgumentParser(description="Audio -> DOCX с таймкодами (Whisper)")
	parser.add_argument("audio", help="Путь к аудио файлу")
	parser.add_argument("out", nargs="?", help="Путь к выходному DOCX")
	parser.add_argument("--model", default="small", help="Модель Whisper: tiny/base/small/medium/large")
	parser.add_argument("--models-dir", default=os.environ.get("WHISPER_MODELS_DIR"), help="Каталог с моделями (small.pt и т.д.)")
	parser.add_argument("--ffmpeg", help="Полный путь к ffmpeg.exe (если не в PATH)")
	args = parser.parse_args()

	audio = args.audio
	if not os.path.exists(audio):
		print(f"Файл не найден: {audio}")
		sys.exit(1)

	out = args.out if args.out else os.path.splitext(audio)[0] + ".docx"
	models_dir = str(Path(args.models_dir).expanduser()) if args.models_dir else None

	# Если указан путь к ffmpeg.exe – добавим его папку в PATH на время процесса
	if args.ffmpeg:
		ffmpeg_path = str(Path(args.ffmpeg))
		ffmpeg_dir = str(Path(ffmpeg_path).parent)
		os.environ["PATH"] = ffmpeg_dir + os.pathsep + os.environ.get("PATH", "")
		print(f"Добавлен в PATH: {ffmpeg_dir}")

	print(f"Входной файл: {audio}")
	print(f"Выходной файл: {out}")
	print(f"Модель: {args.model}")
	if models_dir:
		print(f"Каталог моделей: {models_dir}")

	success = transcribe_to_docx(audio, out, model_name=args.model, models_dir=models_dir)
	if not success:
		sys.exit(1)

